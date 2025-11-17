"""
DOCX to text converter.
"""
from docx import Document
from .base import BaseConverter


class DOCXConverter(BaseConverter):
    """Converter for DOCX documents."""

    def convert(self, input_path: str, output_path: str) -> bool:
        """
        Convert DOCX to text.

        Args:
            input_path: Path to DOCX file
            output_path: Path to output text file

        Returns:
            True if successful, False otherwise
        """
        try:
            text = self._extract_text(input_path)
            if text:
                return self._write_output(text, output_path)
            return False
        except Exception as e:
            self.logger.error(f"Error converting DOCX {input_path}: {str(e)}")
            return False

    def _extract_text(self, input_path: str) -> str:
        """
        Extract text from DOCX.

        Args:
            input_path: Path to DOCX file

        Returns:
            Extracted text
        """
        doc = Document(input_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if not text:
                continue

            if self.preserve_structure:
                # Detect headings by style
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    if level.isdigit():
                        prefix = '#' * int(level)
                        text_parts.append(f"\n{prefix} {text}\n")
                    else:
                        text_parts.append(f"\n## {text}\n")
                else:
                    text_parts.append(f"{text}\n")
            else:
                text_parts.append(f"{text}\n")

        # Extract text from tables
        for table in doc.tables:
            if self.preserve_structure:
                text_parts.append("\n[TABLE]\n")

            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)

                if row_text:
                    if self.preserve_structure:
                        text_parts.append(" | ".join(row_text) + "\n")
                    else:
                        text_parts.append(" ".join(row_text) + "\n")

            if self.preserve_structure:
                text_parts.append("[/TABLE]\n\n")

        return ''.join(text_parts).strip()

    @classmethod
    def get_supported_extensions(cls) -> list[str]:
        """Get supported file extensions."""
        return ['.docx']
